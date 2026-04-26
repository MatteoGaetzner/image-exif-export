import hashlib
import sys
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from PIL import Image
from PIL.ExifTags import GPSTAGS, TAGS
from pillow_heif import register_heif_opener

register_heif_opener()

IMAGE_FORMATS = {
    ".jpg": {"datetime": True, "gps": True},
    ".jpeg": {"datetime": True, "gps": True},
    ".tiff": {"datetime": True, "gps": True},
    ".tif": {"datetime": True, "gps": True},
    ".heic": {"datetime": True, "gps": True},
    ".webp": {"datetime": True, "gps": True},
    ".png": {"datetime": False, "gps": False},
    ".bmp": {"datetime": False, "gps": False},
    ".gif": {"datetime": False, "gps": False},
}


def file_hash(filepath: Path) -> str:
    h = hashlib.md5()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def empty_row(filepath: Path) -> dict:
    return {
        "filename": filepath.name,
        "filepath": str(filepath),
        "extension": filepath.suffix.lower(),
        "size_bytes": filepath.stat().st_size,
        "file_hash": None,
        "width": None,
        "height": None,
        "datetime": None,
        "latitude": None,
        "longitude": None,
        "make": None,
        "model": None,
        "error": None,
    }


def extract_gps_exif(gps_info: dict):
    def to_decimal(values, ref):
        d, m, s = values
        decimal = float(d) + float(m) / 60 + float(s) / 3600
        return -decimal if ref in ("S", "W") else decimal

    try:
        lat = to_decimal(gps_info.get("GPSLatitude"), gps_info.get("GPSLatitudeRef"))
        lon = to_decimal(gps_info.get("GPSLongitude"), gps_info.get("GPSLongitudeRef"))
        return lat, lon
    except Exception:
        return None, None


def extract_image_metadata(filepath: Path) -> dict:
    row = empty_row(filepath)
    try:
        row["file_hash"] = file_hash(filepath)
        with Image.open(filepath) as img:
            row["width"], row["height"] = img.width, img.height
            exif = img._getexif()
            if not exif:
                return row
            decoded = {TAGS.get(k, k): v for k, v in exif.items()}
            row["datetime"] = decoded.get("DateTimeOriginal") or decoded.get("DateTime")
            row["make"] = decoded.get("Make")
            row["model"] = decoded.get("Model")
            raw_gps = decoded.get("GPSInfo")
            if raw_gps:
                gps_info = {GPSTAGS.get(k, k): v for k, v in raw_gps.items()}
                row["latitude"], row["longitude"] = extract_gps_exif(gps_info)
    except Exception as e:
        row["error"] = str(e)
    return row


def process_directory(directory: str) -> pd.DataFrame:
    dir_path = Path(directory)
    if not dir_path.is_dir():
        raise ValueError(f"Not a valid directory: {directory}")

    rows = []
    for filepath in sorted(dir_path.rglob("*")):
        if filepath.suffix.lower() in IMAGE_FORMATS:
            rows.append(extract_image_metadata(filepath))

    if not rows:
        print("No supported image files found.")
        return pd.DataFrame()

    df = pd.DataFrame(rows)

    # Deduplicate by file hash, keeping the first occurrence (alphabetically sorted path)
    before = len(df)
    df = df.drop_duplicates(subset="file_hash", keep="first")
    dupes = before - len(df)
    if dupes:
        print(f"Removed {dupes} duplicate image(s).")

    return df.reset_index(drop=True)


def parse_exif_datetime(dt_str: str | None) -> datetime | None:
    if not dt_str:
        return None
    try:
        return datetime.strptime(dt_str, "%Y:%m:%d %H:%M:%S")
    except Exception:
        return None


def format_caption(row) -> str:
    dt = parse_exif_datetime(row.get("datetime"))
    date_str = dt.strftime("%Y-%m-%d") if dt else "DATE MISSING"
    time_str = dt.strftime("%H:%M:%S") if dt else "TIME MISSING"
    caption = f"{date_str}  {time_str}"
    if dt is None:
        caption += f"  |  {row.get('filename', '')}"
    return caption


def _add_image_block(doc: Document, row, usable_width: int):
    from docx.shared import Pt

    filepath = Path(row["filepath"])

    try:
        with Image.open(filepath) as img:
            w_px, h_px = img.width, img.height
        h = int(usable_width * h_px / w_px)
        doc.add_picture(str(filepath), width=usable_width, height=h)
    except Exception as e:
        doc.add_paragraph(f"[Could not embed {filepath.name}: {e}]")

    # keep_with_next prevents a page break between the image and its caption
    img_para = doc.paragraphs[-1]
    img_para.paragraph_format.keep_with_next = True
    img_para.paragraph_format.space_before = Pt(0)
    img_para.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph(format_caption(row))
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)


def generate_word_doc(df: pd.DataFrame, output_path: str):
    from docx.shared import Pt

    df = df.copy()
    df["_dt_parsed"] = df["datetime"].apply(parse_exif_datetime)
    df = df.sort_values("_dt_parsed", na_position="last").reset_index(drop=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    section = doc.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin

    rows = [row for _, row in df.iterrows() if Path(row["filepath"]).exists()]

    for row in rows:
        _add_image_block(doc, row, usable_width)

    doc.save(output_path)
    print(f"Saved Word document → {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python main.py <path/to/image/directory> [output_base_name]")
        sys.exit(1)

    input_dir = sys.argv[1]
    base_name = sys.argv[2] if len(sys.argv) > 2 else "image_metadata"
    base_name = base_name.removesuffix(".csv").removesuffix(".docx")

    df = process_directory(input_dir)
    if not df.empty:
        df.to_csv(f"{base_name}.csv", index=False)
        print(f"Saved {len(df)} records → {base_name}.csv")

        generate_word_doc(df, f"{base_name}.docx")

        print(df[["filename", "datetime", "latitude", "longitude"]].to_string(index=False))
